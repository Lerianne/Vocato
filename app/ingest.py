"""Ingest career documents into the local vector store.

Sources:
  ~/Downloads/Resumes + CL   — resumes, cover letters, job descriptions
  memory/coffee-chats/       — networking conversation notes
  memory/sessions/           — past coaching session summaries

Incremental: files are hashed; only new/changed files are re-embedded.
Usage: python ingest.py [--force]
"""

import hashlib
import re
import sys
from pathlib import Path

from store import VectorStore

DOCS_DIR = Path.home() / "Downloads" / "Resumes + CL"
MEMORY_DIR = Path(__file__).parent / "memory"
SKIP_EXT = {
    ".jpg", ".jpeg", ".png", ".webp", ".heic", ".gif",
    ".mp4", ".mov", ".zip", ".ds_store", ".sample",
}
CHUNK_CHARS = 3000  # ~750 tokens
OVERLAP = 300


# -- parsers -----------------------------------------------------------------

def parse_pdf(path: Path) -> str:
    from pypdf import PdfReader
    try:
        return "\n".join(p.extract_text() or "" for p in PdfReader(path).pages)
    except Exception as e:
        print(f"  ! pdf error {path.name}: {e}")
        return ""


def parse_docx(path: Path) -> str:
    import docx
    try:
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"  ! docx error {path.name}: {e}")
        return ""


def parse_tabular(path: Path) -> str:
    import pandas as pd
    try:
        if path.suffix.lower() == ".csv":
            df = pd.read_csv(path)
            return f"CSV file {path.name}:\n{df.to_string(max_rows=200)}"
        sheets = pd.read_excel(path, sheet_name=None)
        out = []
        for name, df in sheets.items():
            out.append(f"Sheet '{name}':\n{df.to_string(max_rows=200)}")
        return f"Spreadsheet {path.name}:\n" + "\n\n".join(out)
    except Exception as e:
        print(f"  ! tabular error {path.name}: {e}")
        return ""


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".csv": parse_tabular,
    ".xlsx": parse_tabular,
    ".md": lambda p: p.read_text(errors="replace"),
    ".txt": lambda p: p.read_text(errors="replace"),
}


# -- classification ------------------------------------------------------------

def classify(path: Path, text: str) -> str:
    name = path.name.lower()
    rel = str(path).lower()
    if "coffee-chats" in rel:
        return "coffee-chat"
    if "sessions" in rel:
        return "session"

    head = text[:1500].lower()

    # This is the user's own application archive, so a document is theirs by default.
    # Exclude only clear third-party files (e.g. a relative's CV named "Nathan_*").
    # Legal/admin docs are filtered separately below so they never become résumés.
    is_users = "nathan" not in name and "nathan" not in head

    # Legal / administrative / non-career docs — never experience. Route to "other"
    # before anything else so generic words in them can't trip a wrong bucket.
    if any(w in name for w in ("shareholder", "share_certificate", "share_transfer",
                               "incorporation", "amendment", "offer letter",
                               "interview prep", "recommendation letter")):
        return "other"

    # Cover letters — filename hints (incl. French "lettre"/"motivation", and the
    # misspelled "motvation" seen in the archive) plus English + French content cues.
    if re.search(r"\b(cl|cover.?letters?|motivation|lettre)\b", name) or "motvation" in name:
        return "cover-letter"
    if any(w in head for w in ("dear hiring", "dear recruiter", "dear sir",
                               "i am writing to", "i am excited to apply", "sincerely,",
                               "madame", "monsieur", "lettre de motivation",
                               "veuillez agréer", "cordialement", "ma candidature",
                               "je vous écris", "je me permets de")):
        return "cover-letter"

    # The user's résumé. With no saved job postings, a career document that
    # is theirs and isn't a cover letter is a (often role-tailored) résumé. Accept it
    # on a structure cue OR a résumé-ish filename — text extraction sometimes drops
    # the name/header, so we must not require it.
    if is_users and (
        ("resume" in name or "résumé" in name or "cv" in name)
        or any(w in head for w in ("professional summary", "work experience",
                                   "education", "skills", "experience",
                                   "summary", "projects"))
    ):
        return "resume"

    # Job posting / description — kept for future use, but requires a STRONG,
    # posting-specific phrase (not generic words that also appear in legal/other
    # docs). Most personal archives contain no postings, so this rarely fires.
    if any(p in head for p in ("we are looking for", "about the role", "what you'll do",
                               "what you’ll do", "who you are", "what we offer",
                               "job description", "minimum qualifications",
                               "nous recherchons", "profil recherché", "vos missions")):
        return "job-description"

    # Unknown / non-career → stay SAFE: never assume it's their experience.
    return "other"


def chunk_text(text: str) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    chunks, start = [], 0
    while start < len(text):
        end = start + CHUNK_CHARS
        if end < len(text):  # try to break at a paragraph/sentence boundary
            window = text[start:end]
            cut = max(window.rfind("\n\n"), window.rfind(". "))
            if cut > CHUNK_CHARS // 2:
                end = start + cut + 1
        chunks.append(text[start:end].strip())
        start = end - OVERLAP if end < len(text) else end
    return [c for c in chunks if len(c) > 80]


# -- pipeline ------------------------------------------------------------------

def iter_files():
    for base in (DOCS_DIR, MEMORY_DIR / "coffee-chats", MEMORY_DIR / "sessions"):
        if not base.exists():
            continue
        for path in sorted(base.rglob("*")):
            if not path.is_file() or path.name.startswith((".", "~$")):
                continue
            if ".git" in path.parts or path.suffix.lower() in SKIP_EXT:
                continue
            if path.suffix.lower() in PARSERS:
                yield path


def file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main(force: bool = False):
    store = VectorStore()
    if force:
        store.hashes = {}
    new = skipped = failed = 0
    for path in iter_files():
        key = str(path)
        h = file_hash(path)
        if store.hashes.get(key) == h:
            skipped += 1
            continue
        text = PARSERS[path.suffix.lower()](path)
        chunks = chunk_text(text)
        if not chunks:
            failed += 1
            store.hashes[key] = h  # don't retry unparseable files every run
            continue
        doc_type = classify(path, text)
        store.remove_source(key)
        store.add([
            {
                "text": c,
                "source": key,
                "name": path.name,
                "doc_type": doc_type,
                "part": i,
            }
            for i, c in enumerate(chunks)
        ])
        store.hashes[key] = h
        new += 1
        print(f"  + [{doc_type:>14}] {path.name} ({len(chunks)} chunks)")
        if new % 20 == 0:
            store.save()  # checkpoint
    store.save()
    print(f"\nDone: {new} ingested, {skipped} unchanged, {failed} empty/unparseable.")
    print(f"Index: {len(store.chunks)} chunks total.")


def audit() -> None:
    """Print the current index grouped by classified type, so misfiles are easy
    to spot. The 'other' and 'job-description' groups are the ones to eyeball."""
    from collections import Counter
    store = VectorStore()
    by_type: dict[str, set] = {}
    for c in store.chunks:
        by_type.setdefault(c["doc_type"], set()).add(c["name"])
    print(f"Index: {len(store.chunks)} chunks, "
          f"{sum(len(v) for v in by_type.values())} files\n")
    for dtype in ("resume", "cover-letter", "job-description", "other",
                  "coffee-chat", "session"):
        files = sorted(by_type.get(dtype, ()))
        if not files:
            continue
        print(f"===== {dtype}: {len(files)} files =====")
        for n in files:
            print(f"  {n}")
        print()


if __name__ == "__main__":
    if "--audit" in sys.argv:
        audit()
    else:
        main(force="--force" in sys.argv)
